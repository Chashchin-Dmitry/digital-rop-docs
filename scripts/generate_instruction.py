#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор инструкции "Цифровой РОП - Облачная версия"

Для SaaS-версии платформы https://rop.bvmax.ru
Разработчик: BVMax (https://bvmax.ru)

Возможности:
- Поэтапная генерация (по разделам)
- Пропуск уже созданных частей
- Принудительная перезапись разделов
- Профессиональное форматирование Word
- Автоматическое оглавление
- Корпоративные стили BVMax
"""

import os
import sys
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import argparse
from datetime import datetime
import re


class InstructionGenerator:
    def __init__(self, base_path):
        self.base_path = Path(base_path)
        self.parts_path = self.base_path / "Части_инструкции"
        self.screenshots_path = self.base_path / "СКРИНШОТЫ"
        self.output_path = self.base_path / "Финальная_инструкция.docx"
        self.progress_file = self.base_path / "scripts" / "progress.json"
        self.screenshot_mapping_file = self.base_path / "scripts" / "screenshot_mapping.json"

        # Загружаем маппинг скриншотов
        self.screenshot_mapping = self.load_screenshot_mapping()

        # Загружаем маппинг гиперссылок
        self.hyperlink_mapping_file = self.base_path / "hyperlink_mapping.json"
        self.hyperlink_mapping = self.load_hyperlink_mapping()

        # Структура финальной инструкции для ОБЛАЧНОЙ версии
        self.sections = {
            "intro": {
                "title": "РАЗДЕЛ I: БЫСТРЫЙ СТАРТ",
                "parts": [
                    "00_Что_такое_Цифровой_РОП",
                    "00_Регистрация_и_вход",
                    "00_Первые_шаги"
                ]
            },
            "settings": {
                "title": "РАЗДЕЛ II: НАСТРОЙКИ",
                "parts": [
                    "01_Настройки_Шаблоны_скриптов",
                    "01_Настройки_Скрипты_и_промты",
                    "01_Настройки_Конверсия",
                    "01_Настройки_Дополнительные_промты",
                    "01_Настройки_Таблицы",
                    "01_Настройки_Пользователи"
                ]
            },
            "analytics": {
                "title": "РАЗДЕЛ III: АНАЛИТИКА",
                "parts": [
                    "02_Аналитика_Коммуникации",
                    "02_Аналитика_История_сделок",
                    "02_Аналитика_Менеджеры",
                    "02_Аналитика_Таблицы"
                ]
            },
            "charts": {
                "title": "РАЗДЕЛ IV: ГРАФИКИ",
                "parts": [
                    "03_Графики_Оценка",
                    "03_Графики_Менеджеры"
                ]
            },
            "tests": {
                "title": "РАЗДЕЛ V: ТЕСТЫ",
                "parts": [
                    "04_Тесты"
                ]
            },
            "billing": {
                "title": "РАЗДЕЛ VI: БИЛЛИНГ",
                "parts": [
                    "05_Биллинг"
                ]
            },
            "faq": {
                "title": "РАЗДЕЛ VII: FAQ",
                "parts": [
                    "06_FAQ"
                ]
            },
            "glossary": {
                "title": "РАЗДЕЛ VIII: СЛОВАРЬ ТЕРМИНОВ",
                "parts": [
                    "07_Словарь_терминов"
                ]
            }
        }

        self.load_progress()

    def load_screenshot_mapping(self):
        """Загружает маппинг скриншотов из JSON файла"""
        if self.screenshot_mapping_file.exists():
            with open(self.screenshot_mapping_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return data.get('screenshot_mapping', {})
        return {}

    def load_hyperlink_mapping(self):
        """Загружает маппинг гиперссылок из JSON файла"""
        if self.hyperlink_mapping_file.exists():
            with open(self.hyperlink_mapping_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return data.get('hyperlink_mapping', {})
        return {}

    def load_progress(self):
        """Загружает прогресс генерации"""
        if self.progress_file.exists():
            with open(self.progress_file, 'r', encoding='utf-8') as f:
                self.progress = json.load(f)
        else:
            self.progress = {"completed_sections": [], "completed_parts": []}

    def save_progress(self):
        """Сохраняет прогресс генерации"""
        self.progress_file.parent.mkdir(exist_ok=True)
        with open(self.progress_file, 'w', encoding='utf-8') as f:
            json.dump(self.progress, f, ensure_ascii=False, indent=2)

    def setup_document_styles(self, doc):
        """Настройка профессиональных стилей документа"""

        # Настройка базового шрифта документа
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15
        normal_style.paragraph_format.widow_control = True

        # Корпоративный заголовок документа (зелёный как в интерфейсе)
        if 'Corporate Title' not in [s.name for s in doc.styles]:
            title_style = doc.styles.add_style('Corporate Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(28)
            title_font.bold = True
            title_font.color.rgb = RGBColor(46, 204, 113)  # Зелёный #2ECC71 (как в интерфейсе)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(30)
            title_style.paragraph_format.space_before = Pt(0)

        # Подзаголовок документа
        if 'Corporate Subtitle' not in [s.name for s in doc.styles]:
            subtitle_style = doc.styles.add_style('Corporate Subtitle', WD_STYLE_TYPE.PARAGRAPH)
            subtitle_font = subtitle_style.font
            subtitle_font.name = 'Calibri'
            subtitle_font.size = Pt(18)
            subtitle_font.italic = True
            subtitle_font.color.rgb = RGBColor(89, 89, 89)  # Серый
            subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_style.paragraph_format.space_after = Pt(24)

        # Заголовки разделов (Heading 1) - зелёный как в интерфейсе
        heading1_style = doc.styles['Heading 1']
        heading1_font = heading1_style.font
        heading1_font.name = 'Calibri'
        heading1_font.size = Pt(20)
        heading1_font.bold = True
        heading1_font.color.rgb = RGBColor(46, 204, 113)  # Зелёный #2ECC71
        heading1_style.paragraph_format.space_before = Pt(24)
        heading1_style.paragraph_format.space_after = Pt(12)
        heading1_style.paragraph_format.keep_with_next = True

        # Заголовки подразделов (Heading 2)
        heading2_style = doc.styles['Heading 2']
        heading2_font = heading2_style.font
        heading2_font.name = 'Calibri'
        heading2_font.size = Pt(16)
        heading2_font.bold = True
        heading2_font.color.rgb = RGBColor(40, 40, 40)  # Темно-серый
        heading2_style.paragraph_format.space_before = Pt(18)
        heading2_style.paragraph_format.space_after = Pt(8)
        heading2_style.paragraph_format.keep_with_next = True

        # Заголовки пунктов (Heading 3)
        heading3_style = doc.styles['Heading 3']
        heading3_font = heading3_style.font
        heading3_font.name = 'Calibri'
        heading3_font.size = Pt(14)
        heading3_font.bold = True
        heading3_font.color.rgb = RGBColor(122, 122, 122)  # Серый
        heading3_style.paragraph_format.space_before = Pt(12)
        heading3_style.paragraph_format.space_after = Pt(6)

        # Стиль для скриншотов
        if 'Screenshot Reference' not in [s.name for s in doc.styles]:
            screenshot_style = doc.styles.add_style('Screenshot Reference', WD_STYLE_TYPE.PARAGRAPH)
            screenshot_font = screenshot_style.font
            screenshot_font.name = 'Calibri'
            screenshot_font.size = Pt(10)
            screenshot_font.bold = True
            screenshot_font.italic = True
            screenshot_font.color.rgb = RGBColor(112, 173, 71)  # Зеленый
            screenshot_style.paragraph_format.space_before = Pt(6)
            screenshot_style.paragraph_format.space_after = Pt(3)
            screenshot_style.paragraph_format.left_indent = Inches(0.25)

        # Стиль для важных примечаний
        if 'Important Note' not in [s.name for s in doc.styles]:
            note_style = doc.styles.add_style('Important Note', WD_STYLE_TYPE.PARAGRAPH)
            note_font = note_style.font
            note_font.name = 'Calibri'
            note_font.size = Pt(11)
            note_font.bold = False
            note_font.italic = True
            note_font.color.rgb = RGBColor(120, 120, 120)
            note_style.paragraph_format.space_before = Pt(8)
            note_style.paragraph_format.space_after = Pt(8)
            note_style.paragraph_format.left_indent = Inches(0.5)
            note_style.paragraph_format.right_indent = Inches(0.25)

        # Стиль для кода/URL
        if 'Code Text' not in [s.name for s in doc.styles]:
            code_style = doc.styles.add_style('Code Text', WD_STYLE_TYPE.CHARACTER)
            code_font = code_style.font
            code_font.name = 'Consolas'
            code_font.size = Pt(10)
            code_font.color.rgb = RGBColor(68, 68, 68)

        # Стиль для блоков описания интерфейса
        if 'Interface Block' not in [s.name for s in doc.styles]:
            interface_style = doc.styles.add_style('Interface Block', WD_STYLE_TYPE.PARAGRAPH)
            interface_font = interface_style.font
            interface_font.name = 'Calibri'
            interface_font.size = Pt(11)
            interface_font.color.rgb = RGBColor(89, 89, 89)
            interface_style.paragraph_format.left_indent = Inches(0.4)
            interface_style.paragraph_format.space_before = Pt(6)
            interface_style.paragraph_format.space_after = Pt(6)

        # Стиль для технических деталей
        if 'Technical Block' not in [s.name for s in doc.styles]:
            tech_style = doc.styles.add_style('Technical Block', WD_STYLE_TYPE.PARAGRAPH)
            tech_font = tech_style.font
            tech_font.name = 'Calibri'
            tech_font.size = Pt(11)
            tech_font.color.rgb = RGBColor(52, 73, 94)
            tech_style.paragraph_format.left_indent = Inches(0.4)
            tech_style.paragraph_format.space_before = Pt(6)
            tech_style.paragraph_format.space_after = Pt(6)

        # Стиль для списков
        list_style = doc.styles['List Bullet']
        list_style.paragraph_format.left_indent = Inches(0.25)
        list_style.paragraph_format.space_after = Pt(3)

    def create_document(self):
        """Создание нового документа с корпоративным оформлением BVMax"""
        doc = Document()

        # Настройка полей страницы A4
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

        # Настройка стилей документа
        self.setup_document_styles(doc)

        # Добавляем логотип BVMax
        logo_path = self.base_path / "LOGO" / "bvmax_logo.png"
        if logo_path.exists():
            logo_paragraph = doc.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_paragraph.add_run()
            run.add_picture(str(logo_path), width=Inches(2.5))
            doc.add_paragraph()  # Отступ после лого

        # Корпоративная титульная страница
        title = doc.add_paragraph("Цифровой РОП", style='Corporate Title')

        subtitle = doc.add_paragraph("Полное руководство пользователя", style='Corporate Subtitle')

        # Добавляем пространство
        doc.add_paragraph()
        doc.add_paragraph()

        # Информационный блок
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Заполняем информационную таблицу
        cells = info_table.rows[0].cells
        cells[0].text = "Система:"
        cells[1].text = "Цифровой РОП (облачная версия)"

        cells = info_table.rows[1].cells
        cells[0].text = "Разработчик:"
        cells[1].text = "BVMax (https://bvmax.ru)"

        cells = info_table.rows[2].cells
        cells[0].text = "Дата создания:"
        cells[1].text = datetime.now().strftime('%d.%m.%Y')

        cells = info_table.rows[3].cells
        cells[0].text = "URL системы:"
        cells[1].text = "https://rop.bvmax.ru/login"

        # Форматирование таблицы
        for row in info_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                if cell == row.cells[0]:
                    cell.paragraphs[0].runs[0].font.bold = True

        # Добавляем пространство и описание
        doc.add_paragraph()
        doc.add_paragraph()

        description = doc.add_paragraph(
            "Данная инструкция содержит подробное описание всех функций облачной системы "
            "анализа телефонных переговоров с использованием искусственного интеллекта. "
            "Документ предназначен для менеджеров, руководителей отделов продаж и "
            "администраторов компаний-клиентов платформы Цифровой РОП."
        )
        description.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Контакты поддержки
        doc.add_paragraph()
        support = doc.add_paragraph("Контакты поддержки:")
        support.runs[0].font.bold = True
        doc.add_paragraph("Telegram: @tchashchin")
        doc.add_paragraph("Телефон: +79670047879")

        # Разрыв страницы
        doc.add_page_break()

        return doc

    def add_table_of_contents(self, doc):
        """Добавление оглавления"""
        toc_header = doc.add_paragraph("ОГЛАВЛЕНИЕ", style='Heading 1')
        toc_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Инструкция для автоматического оглавления
        instruction = doc.add_paragraph(
            "Для автоматического создания оглавления в Microsoft Word:\n"
            "1. Установите курсор в это место\n"
            "2. Перейдите в меню 'Ссылки' → 'Оглавление'\n"
            "3. Выберите 'Автособираемое оглавление 1'\n"
            "4. Удалите этот текст после создания оглавления"
        )
        instruction.style = 'Important Note'

        # Ручное оглавление
        doc.add_paragraph()

        toc_content = [
            ("РАЗДЕЛ I: БЫСТРЫЙ СТАРТ", "Heading 2"),
            ("    1. Что такое Цифровой РОП", None),
            ("    2. Регистрация и вход в систему", None),
            ("    3. Первые шаги", None),
            ("", None),
            ("РАЗДЕЛ II: НАСТРОЙКИ", "Heading 2"),
            ("    4. Шаблоны скриптов (чек-листы)", None),
            ("    5. Скрипты и промты", None),
            ("    6. Настройка конверсии", None),
            ("    7. Дополнительные промты", None),
            ("    8. Таблицы (кастомная аналитика)", None),
            ("    9. Пользователи", None),
            ("", None),
            ("РАЗДЕЛ III: АНАЛИТИКА", "Heading 2"),
            ("    10. Коммуникации (все звонки)", None),
            ("    11. История сделок", None),
            ("    12. Менеджеры (рейтинги)", None),
            ("    13. Таблицы (результаты)", None),
            ("", None),
            ("РАЗДЕЛ IV: ГРАФИКИ", "Heading 2"),
            ("    14. Оценка качества", None),
            ("    15. Динамика менеджеров", None),
            ("", None),
            ("РАЗДЕЛ V: ТЕСТЫ", "Heading 2"),
            ("    16. Автоматическое тестирование", None),
            ("", None),
            ("РАЗДЕЛ VI: БИЛЛИНГ", "Heading 2"),
            ("    17. Статистика использования", None),
            ("", None),
            ("РАЗДЕЛ VII: FAQ", "Heading 2"),
            ("    18. Часто задаваемые вопросы", None),
            ("", None),
            ("РАЗДЕЛ VIII: СЛОВАРЬ ТЕРМИНОВ", "Heading 2"),
            ("    19. Глоссарий", None),
        ]

        for item_text, style in toc_content:
            if not item_text:
                doc.add_paragraph()
                continue

            if style:
                p = doc.add_paragraph(item_text, style=style)
            else:
                p = doc.add_paragraph(item_text)
                p.paragraph_format.left_indent = Inches(0.25)

        doc.add_page_break()

    def find_screenshot(self, filename):
        """Поиск скриншота используя JSON маппинг и fallback к папкам"""
        filename = filename.strip()

        # Убираем все звездочки из имени файла
        filename = filename.replace('*', '').strip()

        # Если нет расширения, добавляем .png
        if not filename.endswith('.png'):
            filename += '.png'

        # Сначала проверяем в JSON маппинге
        if filename in self.screenshot_mapping:
            screenshot_path = self.base_path / self.screenshot_mapping[filename]
            if screenshot_path.exists():
                return screenshot_path

        # Fallback: поиск по частичному совпадению в маппинге
        filename_base = filename.lower().replace('.png', '').replace('.jpg', '')
        for mapped_name, mapped_path in self.screenshot_mapping.items():
            mapped_base = mapped_name.lower().replace('.png', '').replace('.jpg', '')
            if filename_base in mapped_base or mapped_base in filename_base:
                screenshot_path = self.base_path / mapped_path
                if screenshot_path.exists():
                    return screenshot_path

        # Fallback: поиск во всех подпапках
        for root, dirs, files in os.walk(self.screenshots_path):
            for file in files:
                if file.lower() == filename.lower():
                    return Path(root) / file

                file_base = file.lower().replace('.png', '').replace('.jpg', '')
                filename_base = filename.lower().replace('.png', '').replace('.jpg', '')

                if filename_base in file_base or file_base in filename_base:
                    return Path(root) / file
        return None

    def process_text_formatting(self, text):
        """Очищает markdown символы"""
        text = re.sub(r'[#`]', '', text).strip()
        return text

    def add_formatted_paragraph(self, doc, text, style=None):
        """Добавляет параграф с форматированием жирного текста и гиперссылок"""
        if style:
            p = doc.add_paragraph(style=style)
        else:
            p = doc.add_paragraph()

        if '[' in text and '](#' in text:
            self._add_complex_formatting(p, text)
        else:
            parts = re.split(r'(\*\*[^*]+\*\*)', text)

            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    bold_text = part[2:-2]
                    run = p.add_run(bold_text)
                    run.font.bold = True
                elif part:
                    clean_part = re.sub(r'[*#`]', '', part)
                    if clean_part.strip():
                        run = p.add_run(clean_part)

        return p

    def _add_complex_formatting(self, paragraph, text):
        """Добавляет текст с гиперссылками и жирным форматированием"""
        hyperlink_pattern = r'\[([^\]]+)\]\(#([^)]+)\)'
        parts = re.split(hyperlink_pattern, text)

        for i, part in enumerate(parts):
            if i % 3 == 0:
                if part:
                    self._add_bold_formatting(paragraph, part)
            elif i % 3 == 1:
                link_text = part
            elif i % 3 == 2:
                anchor = "#" + part
                target_heading = self.hyperlink_mapping.get(anchor, None)

                if target_heading:
                    self._create_internal_hyperlink(paragraph, link_text, target_heading)
                else:
                    hyperlink = paragraph.add_run(link_text)
                    hyperlink.font.color.rgb = RGBColor(255, 0, 0)

    def _add_bold_formatting(self, paragraph, text):
        """Добавляет текст с жирным форматированием"""
        parts = re.split(r'(\*\*[^*]+\*\*)', text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                bold_text = part[2:-2]
                run = paragraph.add_run(bold_text)
                run.font.bold = True
            elif part:
                clean_part = re.sub(r'[*#`]', '', part)
                if clean_part.strip():
                    run = paragraph.add_run(clean_part)

    def add_interface_block(self, doc, text):
        """Добавляет блок описания интерфейса с синей границей"""
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        borders_xml = '''
        <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:left w:val="single" w:sz="24" w:color="3498DB"/>
            <w:top w:val="nil"/>
            <w:right w:val="nil"/>
            <w:bottom w:val="nil"/>
        </w:tcBorders>'''

        shading_xml = '''
        <w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
               w:val="clear" w:color="auto" w:fill="EBF3FD"/>'''

        try:
            borders_element = parse_xml(borders_xml)
            shading_element = parse_xml(shading_xml)

            tc_pr = cell._tc.get_or_add_tcPr()
            tc_pr.append(borders_element)
            tc_pr.append(shading_element)
        except Exception as e:
            print(f"DEBUG: Border/shading failed: {e}")

        cell_para = cell.paragraphs[0]
        self._add_formatted_text_to_paragraph(cell_para, text)

        cell_para.paragraph_format.space_before = Pt(6)
        cell_para.paragraph_format.space_after = Pt(6)

        for run in cell_para.runs:
            run.font.color.rgb = RGBColor(40, 40, 40)
            run.font.italic = True

        return table

    def _add_formatted_text_to_paragraph(self, paragraph, text):
        """Добавляет форматированный текст в параграф"""
        if '[' in text and '](#' in text:
            self._add_complex_formatting(paragraph, text)
        else:
            parts = re.split(r'(\*\*[^*]+\*\*)', text)

            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    bold_text = part[2:-2]
                    run = paragraph.add_run(bold_text)
                    run.font.bold = True
                elif part:
                    clean_part = re.sub(r'[*#`]', '', part)
                    if clean_part.strip():
                        run = paragraph.add_run(clean_part)

    def add_technical_block(self, doc, text):
        """Добавляет технический блок с оранжевой границей"""
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        borders_xml = '''
        <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:left w:val="single" w:sz="24" w:color="E67E22"/>
            <w:top w:val="nil"/>
            <w:right w:val="nil"/>
            <w:bottom w:val="nil"/>
        </w:tcBorders>'''

        shading_xml = '''
        <w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
               w:val="clear" w:color="auto" w:fill="FDF2E9"/>'''

        try:
            borders_element = parse_xml(borders_xml)
            shading_element = parse_xml(shading_xml)

            tc_pr = cell._tc.get_or_add_tcPr()
            tc_pr.append(borders_element)
            tc_pr.append(shading_element)
        except Exception as e:
            print(f"DEBUG: Technical border/shading failed: {e}")

        cell_para = cell.paragraphs[0]
        self._add_formatted_text_to_paragraph(cell_para, text)

        cell_para.paragraph_format.space_before = Pt(6)
        cell_para.paragraph_format.space_after = Pt(6)

        for run in cell_para.runs:
            run.font.color.rgb = RGBColor(80, 80, 80)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)

        return table

    def add_code_block(self, doc, text, title=""):
        """Добавляет блок кода с фоном"""
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'

        cell = table.cell(0, 0)
        cell.width = Inches(6)

        shading_elm = parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ))
        cell._tc.get_or_add_tcPr().append(shading_elm)

        if title:
            title_p = cell.paragraphs[0]
            title_run = title_p.add_run(title)
            title_run.font.bold = True
            title_run.font.size = Pt(10)
            title_run.font.color.rgb = RGBColor(46, 204, 113)  # Зелёный #2ECC71
            cell.add_paragraph()

        code_p = cell.add_paragraph()
        code_run = code_p.add_run(text)
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(9)
        code_run.font.color.rgb = RGBColor(40, 40, 40)

        return table

    def _create_internal_hyperlink(self, paragraph, link_text, target_heading):
        """Создает внутреннюю гиперссылку Word"""
        bookmark_id = target_heading.replace(" ", "_").replace(".", "_").replace("-", "_").replace("(", "").replace(")", "")

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_id)

        new_run = OxmlElement('w:r')

        rPr = OxmlElement('w:rPr')

        c = OxmlElement('w:color')
        c.set(qn('w:val'), '2A6099')
        rPr.append(c)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)

        t = OxmlElement('w:t')
        t.text = link_text
        new_run.append(t)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _add_bookmark(self, paragraph, bookmark_name):
        """Добавляет закладку к параграфу"""
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        r = run._r

        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), '0')
        start.set(qn('w:name'), bookmark_name)
        r.insert(0, start)

        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), '0')
        end.set(qn('w:name'), bookmark_name)
        r.append(end)

    def read_part_content(self, part_name):
        """Чтение содержимого части инструкции"""
        # Сначала проверяем версию _NEW если существует
        new_file = self.parts_path / f"{part_name}_NEW.md"
        if new_file.exists():
            with open(new_file, 'r', encoding='utf-8') as f:
                return f.read()

        # Иначе используем обычный файл
        part_file = self.parts_path / f"{part_name}.md"
        if part_file.exists():
            with open(part_file, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            return f"Данный раздел находится в разработке и будет добавлен в следующих версиях."

    def add_markdown_table(self, doc, table_lines):
        """Добавляет таблицу Markdown в документ Word"""
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        valid_lines = []
        for line in table_lines:
            line = line.strip()
            if line and line.startswith('|') and not line.replace('|', '').replace('-', '').replace(' ', '').replace(':', '') == '':
                valid_lines.append(line)

        if len(valid_lines) < 2:
            return

        header_line = valid_lines[0]
        headers = [cell.strip() for cell in header_line.split('|')[1:-1]]

        data_lines = []
        for line in valid_lines[1:]:
            if not line.replace('|', '').replace('-', '').replace(' ', '').replace(':', '') == '':
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if len(cells) == len(headers):
                    data_lines.append(cells)

        if not data_lines:
            return

        table = doc.add_table(rows=1 + len(data_lines), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = 'Table Grid'

        header_row = table.rows[0]
        for i, header_text in enumerate(headers):
            cell = header_row.cells[i]
            paragraph = cell.paragraphs[0]
            paragraph.text = header_text
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            run = paragraph.runs[0]
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)

            shading_elm = parse_xml(r'<w:shd {} w:fill="2ECC71"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))  # BVMax зелёный
            cell._tc.get_or_add_tcPr().append(shading_elm)

        for row_idx, row_data in enumerate(data_lines, 1):
            row = table.rows[row_idx]
            for col_idx, cell_text in enumerate(row_data):
                cell = row.cells[col_idx]
                paragraph = cell.paragraphs[0]
                paragraph.text = cell_text

                run = paragraph.runs[0]
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(60, 60, 60)

    def add_section_to_doc(self, doc, section_key, force_regenerate=False):
        """Добавление раздела в документ"""
        section_info = self.sections[section_key]

        if not force_regenerate and section_key in self.progress['completed_sections']:
            print(f"Раздел '{section_info['title']}' уже создан, пропускаем...")
            return

        print(f"Создаём раздел: {section_info['title']}")

        section_heading = doc.add_paragraph(section_info['title'], style='Heading 1')

        section_bookmark_id = section_info['title'].replace(" ", "_").replace(".", "_").replace("-", "_").replace("(", "").replace(")", "")
        self._add_bookmark(section_heading, section_bookmark_id)

        for part_name in section_info['parts']:
            if not force_regenerate and part_name in self.progress['completed_parts']:
                print(f"  Часть '{part_name}' уже создана, пропускаем...")
                continue

            print(f"  Добавляем часть: {part_name}")

            content = self.read_part_content(part_name)

            lines = content.split('\n')
            current_list_level = 0

            in_code_block = False
            code_block_content = []

            for line_num, line in enumerate(lines, 1):
                line = line.rstrip()

                # Специальные блоки
                if line.startswith('{INTERFACE}'):
                    text = line[11:].strip()
                    self.add_interface_block(doc, text)
                    current_list_level = 0
                    continue
                elif line.startswith('{TECHNICAL}'):
                    text = line[11:].strip()
                    self.add_technical_block(doc, text)
                    current_list_level = 0
                    continue

                # Блоки кода
                if line.strip() == '```':
                    if in_code_block:
                        if code_block_content:
                            self.add_code_block(doc, '\n'.join(code_block_content))
                        code_block_content = []
                        in_code_block = False
                    else:
                        in_code_block = True
                    continue

                if in_code_block:
                    code_block_content.append(line)
                    continue

                if not line:
                    doc.add_paragraph()
                    current_list_level = 0
                    continue

                if line.startswith('# '):
                    clean_title = line[2:].strip()
                    heading_para = doc.add_paragraph(clean_title, style='Heading 2')
                    bookmark_id = clean_title.replace(" ", "_").replace(".", "_").replace("-", "_").replace("(", "").replace(")", "")
                    self._add_bookmark(heading_para, bookmark_id)
                    current_list_level = 0

                elif line.startswith('## '):
                    clean_title = line[3:].strip()
                    doc.add_paragraph(clean_title, style='Heading 3')
                    current_list_level = 0

                elif line.startswith('### '):
                    clean_title = line[4:].strip()
                    p = doc.add_paragraph(clean_title)
                    p.runs[0].font.bold = True
                    p.runs[0].font.size = Pt(12)
                    p.runs[0].font.color.rgb = RGBColor(122, 122, 122)
                    current_list_level = 0

                elif line.startswith('|') and '|' in line.strip():
                    table_lines = [line]
                    for next_line_idx in range(line_num, len(lines)):
                        next_line = lines[next_line_idx].strip()
                        if next_line.startswith('#') or (next_line and not next_line.startswith('|')):
                            break
                        if next_line.startswith('|') and '|' in next_line:
                            table_lines.append(next_line)

                    if len(table_lines) >= 2:
                        self.add_markdown_table(doc, table_lines)
                    current_list_level = 0

                elif '.png' in line or 'скриншот' in line.lower() or '[СКРИНШОТ:' in line:
                    if '[СКРИНШОТ:' in line:
                        match = re.search(r'\[СКРИНШОТ:\s*([^\]]+)\]', line)
                        if match:
                            image_filename = match.group(1).strip()
                            screenshot_text = image_filename
                        else:
                            continue
                    else:
                        clean_line = line.replace('*', '').strip()

                        if ' - ' in clean_line:
                            parts = clean_line.split(' - ', 1)
                            image_filename = parts[0].strip()
                            description = parts[1].strip()
                            screenshot_text = f"{image_filename} - {description}"
                        else:
                            image_filename = clean_line.strip()
                            screenshot_text = image_filename

                    image_path = self.find_screenshot(image_filename)
                    if image_path and image_path.exists():
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(str(image_path), width=Inches(6))
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        caption = doc.add_paragraph(screenshot_text, style='Screenshot Reference')
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    current_list_level = 0

                elif line.startswith('- ') or line.startswith('* '):
                    list_text = line[2:].strip()
                    if list_text:
                        p = self.add_formatted_paragraph(doc, list_text, 'List Bullet')
                    current_list_level = 1

                elif re.match(r'^\d+\.\s', line):
                    list_text = re.sub(r'^\d+\.\s', '', line)
                    list_text = re.sub(r'[*#`]', '', list_text).strip()
                    if list_text:
                        p = doc.add_paragraph(list_text, style='List Number')
                    current_list_level = 1

                elif line.startswith('  - ') or line.startswith('    - '):
                    indent_level = (len(line) - len(line.lstrip())) // 2
                    list_text = line.strip()[2:]
                    p = doc.add_paragraph(list_text, style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.25 * (indent_level // 2 + 1))
                    current_list_level = 2

                elif line.startswith('http://') or line.startswith('https://'):
                    self.add_code_block(doc, line, "URL:")
                    current_list_level = 0

                elif '`' in line:
                    clean_line = line.replace('*', '')
                    p = doc.add_paragraph()
                    parts = clean_line.split('`')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            if part:
                                p.add_run(part)
                        else:
                            run = p.add_run(part)
                            run.style = 'Code Text'
                    current_list_level = 0

                elif line.strip():
                    clean_line = self.process_text_formatting(line.strip())
                    if clean_line:
                        if any(keyword in clean_line.lower() for keyword in ['важно', 'внимание', 'примечание', 'note']):
                            p = self.add_formatted_paragraph(doc, line, 'Important Note')
                        else:
                            p = self.add_formatted_paragraph(doc, line)
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    current_list_level = 0

            if part_name not in self.progress['completed_parts']:
                self.progress['completed_parts'].append(part_name)

        if section_key not in self.progress['completed_sections']:
            self.progress['completed_sections'].append(section_key)

        doc.add_page_break()

    def generate(self, sections_to_generate=None, force_regenerate=False):
        """Основная функция генерации"""
        print("=" * 60)
        print("Генератор инструкции Цифровой РОП (облачная версия)")
        print("=" * 60)

        if self.output_path.exists() and not force_regenerate:
            print("Открываем существующий документ...")
            doc = Document(str(self.output_path))
        else:
            print("Создаём новый документ...")
            doc = self.create_document()
            self.add_table_of_contents(doc)
            if force_regenerate:
                self.progress = {"completed_sections": [], "completed_parts": []}

        if sections_to_generate:
            target_sections = sections_to_generate
        else:
            target_sections = list(self.sections.keys())

        for section_key in target_sections:
            if section_key in self.sections:
                self.add_section_to_doc(doc, section_key, force_regenerate)
            else:
                print(f"ОШИБКА: Неизвестный раздел: {section_key}")

        print("Сохраняем документ...")
        doc.save(str(self.output_path))
        self.save_progress()

        print("=" * 60)
        print(f"ГОТОВО: {self.output_path}")
        print(f"Прогресс: {len(self.progress['completed_parts'])} частей создано")
        print("=" * 60)


def main():
    parser = argparse.ArgumentParser(description='Генератор инструкции Цифровой РОП (облачная версия)')
    parser.add_argument('--sections', nargs='+', help='Конкретные разделы для генерации')
    parser.add_argument('--force', action='store_true', help='Принудительная перезапись всех разделов')
    parser.add_argument('--reset', action='store_true', help='Сброс прогресса и создание нового документа')
    parser.add_argument('--list', action='store_true', help='Показать доступные разделы')

    args = parser.parse_args()

    base_path = Path(__file__).parent.parent
    generator = InstructionGenerator(base_path)

    if args.list:
        print("Доступные разделы:")
        for key, info in generator.sections.items():
            print(f"  {key}: {info['title']}")
            for part in info['parts']:
                print(f"      - {part}")
        return

    if args.reset:
        print("Сброс прогресса...")
        if generator.progress_file.exists():
            generator.progress_file.unlink()
        if generator.output_path.exists():
            generator.output_path.unlink()
        generator.load_progress()

    generator.generate(
        sections_to_generate=args.sections,
        force_regenerate=args.force or args.reset
    )


if __name__ == "__main__":
    main()
