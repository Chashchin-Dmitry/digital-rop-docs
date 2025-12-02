#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор финальной инструкции "Цифровой РОП - Полное руководство пользователя"

Возможности:
- Поэтапная генерация (по разделам) 
- Пропуск уже созданных частей
- Принудительная перезапись разделов
- Профессиональное форматирование Word
- Автоматическое оглавление
- Стили корпоративного документооборота
"""

import os
import sys
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import argparse
from datetime import datetime
import re

class InstructionGenerator:
    def __init__(self, base_path):
        self.base_path = Path(base_path)
        self.parts_path = self.base_path / "Части_инструкции"
        self.screenshots_path = self.base_path / "СКРИНШОТЫ"
        self.output_path = self.base_path / "Финальная_инструкция.docx"
        self.progress_file = self.base_path / "scripts" / "progress.json"
        self.screenshot_mapping_file = self.base_path / "scripts" / "screenshot_mapping.json"
        
        # Загружаем маппинг скриншотов
        self.screenshot_mapping = self.load_screenshot_mapping()
        
        # Загружаем маппинг гиперссылок
        self.hyperlink_mapping_file = self.base_path / "hyperlink_mapping.json"
        self.hyperlink_mapping = self.load_hyperlink_mapping()
        
        # Структура финальной инструкции
        self.sections = {
            "intro": {
                "title": "РАЗДЕЛ I: БЫСТРЫЙ СТАРТ",
                "parts": [
                    "00_Что_такое_Цифровой_РОП",
                    "00_Первые_5_минут", 
                    "00_Интерфейс_и_навигация",
                    "00_Роли_и_права_доступа",
                    "00_Вход"
                ]
            },
            "analytics": {
                "title": "РАЗДЕЛ II: АНАЛИТИКА И ОТЧЁТЫ", 
                "parts": [
                    "01_Аналитика_Коммуникации",
                    "01_Аналитика_Менеджеры", 
                    "01_Аналитика_Тесты"
                ]
            },
            "charts": {
                "title": "РАЗДЕЛ III: ГРАФИКИ И ВИЗУАЛИЗАЦИЯ",
                "parts": [
                    "02_Графики_Менеджеры",
                    "02_Графики_Оценка", 
                    "02_Графики_Статус_тестов",
                    "02_Графики_Динамика_тестов"
                ]
            },
            "tables": {
                "title": "РАЗДЕЛ IV: НАСТРОЙКА АНАЛИТИЧЕСКИХ ТАБЛИЦ",
                "parts": [
                    "03_Таблицы",
                    "03_Создание_промтов_пошагово",
                    "03_Готовые_примеры_промтов"
                ]
            },
            "tools": {
                "title": "РАЗДЕЛ V: ДОПОЛНИТЕЛЬНЫЕ ИНСТРУМЕНТЫ", 
                "parts": [
                    "04_Другое"
                ]
            },
            "settings": {
                "title": "РАЗДЕЛ VI: АДМИНИСТРИРОВАНИЕ СИСТЕМЫ",
                "parts": [
                    "07_Администрирование_системы",
                    "05_Настройки_Офисы_и_руководители",
                    "05_Настройки_Подключение", 
                    "05_Настройки_Пользователи",
                    "05_Настройки_Скрипты_и_промты",
                    "05_Настройки_Таблицы"
                ]
            },
            "additional": {
                "title": "РАЗДЕЛ VII: ДОПОЛНИТЕЛЬНЫЕ НАСТРОЙКИ",
                "parts": [
                    "06_Доп_Настройки_Слова_паразиты"
                ]
            },
            "processes": {
                "title": "РАЗДЕЛ VIII: БИЗНЕС-ПРОЦЕССЫ",
                "parts": [
                    "08_Полный_цикл_обработки_звонка",
                    "08_Процесс_создания_тестов",
                    "08_Управление_ролями"
                ]
            },
            "examples": {
                "title": "РАЗДЕЛ IX: ПРАКТИЧЕСКИЕ ПРИМЕРЫ",
                "parts": [
                    "09_Сценарии_для_ролей",
                    "09_Типичные_задачи",
                    "09_FAQ"
                ]
            },
            "technical": {
                "title": "РАЗДЕЛ X: ТЕХНИЧЕСКАЯ ИНФОРМАЦИЯ", 
                "parts": [
                    "10_Словарь_терминов",
                    "12_API_документация_техническая"
                ]
            }
        }
        
        self.load_progress()
    
    def load_screenshot_mapping(self):
        """Загружает маппинг скриншотов из JSON файла"""
        if self.screenshot_mapping_file.exists():
            with open(self.screenshot_mapping_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return data.get('screenshot_mapping', {})
        return {}
    
    def load_hyperlink_mapping(self):
        """Загружает маппинг гиперссылок из JSON файла"""
        if self.hyperlink_mapping_file.exists():
            with open(self.hyperlink_mapping_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return data.get('hyperlink_mapping', {})
        return {}
    
    def load_progress(self):
        """Загружает прогресс генерации"""
        if self.progress_file.exists():
            with open(self.progress_file, 'r', encoding='utf-8') as f:
                self.progress = json.load(f)
        else:
            self.progress = {"completed_sections": [], "completed_parts": []}
    
    def save_progress(self):
        """Сохраняет прогресс генерации"""
        self.progress_file.parent.mkdir(exist_ok=True)
        with open(self.progress_file, 'w', encoding='utf-8') as f:
            json.dump(self.progress, f, ensure_ascii=False, indent=2)
    
    def setup_document_styles(self, doc):
        """Настройка профессиональных стилей документа"""
        
        # Настройка базового шрифта документа
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15
        normal_style.paragraph_format.widow_control = True
        
        # Корпоративный заголовок документа
        if 'Corporate Title' not in [s.name for s in doc.styles]:
            title_style = doc.styles.add_style('Corporate Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(28)
            title_font.bold = True
            title_font.color.rgb = RGBColor(237, 28, 36)  # Корпоративный красный Setl Group
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(30)
            title_style.paragraph_format.space_before = Pt(0)
        
        # Подзаголовок документа
        if 'Corporate Subtitle' not in [s.name for s in doc.styles]:
            subtitle_style = doc.styles.add_style('Corporate Subtitle', WD_STYLE_TYPE.PARAGRAPH)
            subtitle_font = subtitle_style.font
            subtitle_font.name = 'Calibri'
            subtitle_font.size = Pt(18)
            subtitle_font.italic = True
            subtitle_font.color.rgb = RGBColor(89, 89, 89)  # Серый
            subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_style.paragraph_format.space_after = Pt(24)
        
        # Заголовки разделов (Heading 1)
        heading1_style = doc.styles['Heading 1']
        heading1_font = heading1_style.font
        heading1_font.name = 'Calibri'
        heading1_font.size = Pt(20)
        heading1_font.bold = True
        heading1_font.color.rgb = RGBColor(237, 28, 36)  # Setl Group красный
        heading1_style.paragraph_format.space_before = Pt(24)
        heading1_style.paragraph_format.space_after = Pt(12)
        heading1_style.paragraph_format.keep_with_next = True
        
        # Заголовки подразделов (Heading 2)
        heading2_style = doc.styles['Heading 2'] 
        heading2_font = heading2_style.font
        heading2_font.name = 'Calibri'
        heading2_font.size = Pt(16)
        heading2_font.bold = True
        heading2_font.color.rgb = RGBColor(40, 40, 40)  # Setl Group темно-серый
        heading2_style.paragraph_format.space_before = Pt(18)
        heading2_style.paragraph_format.space_after = Pt(8)
        heading2_style.paragraph_format.keep_with_next = True
        
        # Заголовки пунктов (Heading 3)
        heading3_style = doc.styles['Heading 3']
        heading3_font = heading3_style.font  
        heading3_font.name = 'Calibri'
        heading3_font.size = Pt(14)
        heading3_font.bold = True
        heading3_font.color.rgb = RGBColor(122, 122, 122)  # Setl Group серый
        heading3_style.paragraph_format.space_before = Pt(12)
        heading3_style.paragraph_format.space_after = Pt(6)
        
        # Стиль для скриншотов
        if 'Screenshot Reference' not in [s.name for s in doc.styles]:
            screenshot_style = doc.styles.add_style('Screenshot Reference', WD_STYLE_TYPE.PARAGRAPH)
            screenshot_font = screenshot_style.font
            screenshot_font.name = 'Calibri'
            screenshot_font.size = Pt(10)
            screenshot_font.bold = True
            screenshot_font.italic = True
            screenshot_font.color.rgb = RGBColor(112, 173, 71)  # Зеленый
            screenshot_style.paragraph_format.space_before = Pt(6)
            screenshot_style.paragraph_format.space_after = Pt(3)
            screenshot_style.paragraph_format.left_indent = Inches(0.25)
        
        # Стиль для важных примечаний
        if 'Important Note' not in [s.name for s in doc.styles]:
            note_style = doc.styles.add_style('Important Note', WD_STYLE_TYPE.PARAGRAPH)
            note_font = note_style.font
            note_font.name = 'Calibri'
            note_font.size = Pt(11)
            note_font.bold = False  # Убираем жирный
            note_font.italic = True  # Делаем курсивом
            note_font.color.rgb = RGBColor(120, 120, 120)  # Серый цвет
            note_style.paragraph_format.space_before = Pt(8)
            note_style.paragraph_format.space_after = Pt(8)
            note_style.paragraph_format.left_indent = Inches(0.5)  # Больше отступ
            note_style.paragraph_format.right_indent = Inches(0.25)
        
        # Стиль для кода/URL
        if 'Code Text' not in [s.name for s in doc.styles]:
            code_style = doc.styles.add_style('Code Text', WD_STYLE_TYPE.CHARACTER)
            code_font = code_style.font
            code_font.name = 'Consolas'
            code_font.size = Pt(10)
            code_font.color.rgb = RGBColor(68, 68, 68)
        
        # Стиль для блоков описания интерфейса с левой границей
        if 'Interface Block' not in [s.name for s in doc.styles]:
            interface_style = doc.styles.add_style('Interface Block', WD_STYLE_TYPE.PARAGRAPH)
            interface_font = interface_style.font
            interface_font.name = 'Calibri'
            interface_font.size = Pt(11)
            interface_font.color.rgb = RGBColor(89, 89, 89)  # Серый
            interface_style.paragraph_format.left_indent = Inches(0.4)
            interface_style.paragraph_format.space_before = Pt(6)
            interface_style.paragraph_format.space_after = Pt(6)
        
        # Стиль для технических деталей с синей границей
        if 'Technical Block' not in [s.name for s in doc.styles]:
            tech_style = doc.styles.add_style('Technical Block', WD_STYLE_TYPE.PARAGRAPH)
            tech_font = tech_style.font
            tech_font.name = 'Calibri'
            tech_font.size = Pt(11)
            tech_font.color.rgb = RGBColor(52, 73, 94)  # Темно-синий
            tech_style.paragraph_format.left_indent = Inches(0.4)
            tech_style.paragraph_format.space_before = Pt(6)
            tech_style.paragraph_format.space_after = Pt(6)
        
        # Стиль для списков с отступом
        list_style = doc.styles['List Bullet']
        list_style.paragraph_format.left_indent = Inches(0.25)
        list_style.paragraph_format.space_after = Pt(3)
        
        # Настройка таблиц
        table_style = doc.styles['Table Grid']
    
    def create_document(self):
        """Создание нового документа с корпоративным оформлением"""
        doc = Document()
        
        # Настройка полей страницы
        section = doc.sections[0]
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21.0)   # A4
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        
        # Настройка стилей документа
        self.setup_document_styles(doc)
        
        # Добавляем лого компании если есть
        logo_path = self.base_path / "Лого Setl"
        if logo_path.exists():
            logo_files = list(logo_path.glob("*.jpg")) + list(logo_path.glob("*.png"))
            if logo_files:
                logo_paragraph = doc.add_paragraph()
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = logo_paragraph.add_run()
                run.add_picture(str(logo_files[0]), width=Inches(3))
                doc.add_paragraph()  # Отступ после лого
        
        # Корпоративная титульная страница
        title = doc.add_paragraph("Цифровой РОП", style='Corporate Title')
        
        subtitle = doc.add_paragraph("Полное руководство пользователя", style='Corporate Subtitle')
        
        # Добавляем пространство
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Информационный блок
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заполняем информационную таблицу
        cells = info_table.rows[0].cells
        cells[0].text = "Система:"
        cells[1].text = "Цифровой РОП (Цифровой ассистент)"
        
        cells = info_table.rows[1].cells
        cells[0].text = "Компания:"
        cells[1].text = "Setl Group (https://setlgroup.ru)"
        
        cells = info_table.rows[2].cells
        cells[0].text = "Дата создания:"
        cells[1].text = datetime.now().strftime('%d.%m.%Y')
        
        cells = info_table.rows[3].cells
        cells[0].text = "URL системы:"
        cells[1].text = "http://10.28.32.81/"
        
        # Форматирование таблицы
        for row in info_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                if cell == row.cells[0]:  # Первый столбец - жирный
                    cell.paragraphs[0].runs[0].font.bold = True
        
        # Добавляем пространство и описание
        doc.add_paragraph()
        doc.add_paragraph()
        
        description = doc.add_paragraph(
            "Данная инструкция содержит подробное описание всех функций системы "
            "анализа телефонных переговоров с использованием искусственного интеллекта. "
            "Документ предназначен для менеджеров, руководителей офисов и системных "
            "администраторов компании Setl Group."
        )
        description.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Разрыв страницы
        doc.add_page_break()
        
        return doc
    
    def add_table_of_contents(self, doc):
        """Добавление оглавления"""
        toc_header = doc.add_paragraph("ОГЛАВЛЕНИЕ", style='Heading 1')
        toc_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Создаем структурированное оглавление
        doc.add_paragraph()
        
        # Добавляем инструкцию для автоматического оглавления
        instruction = doc.add_paragraph(
            "Для автоматического создания оглавления в Microsoft Word:\n"
            "1. Установите курсор в это место\n"
            "2. Перейдите в меню 'Ссылки' → 'Оглавление'\n"
            "3. Выберите 'Автособираемое оглавление 1'\n"
            "4. Удалите этот текст после создания оглавления"
        )
        instruction.style = 'Important Note'
        
        # Создаем ручное оглавление для предварительного просмотра
        doc.add_paragraph()
        
        toc_content = [
            ("РАЗДЕЛ I: БЫСТРЫЙ СТАРТ", "Heading 2"),
            ("    1. Что такое Цифровой РОП", None),
            ("    2. Первые 5 минут в системе", None),
            ("    3. Интерфейс и навигация", None),
            ("    4. Роли и права доступа", None),
            ("    5. Вход в систему", None),
            ("", None),
            ("РАЗДЕЛ II: АНАЛИТИКА И ОТЧЁТЫ", "Heading 2"),
            ("    6. Коммуникации - просмотр звонков", None),
            ("    7. Менеджеры - рейтинги и производительность", None),
            ("    8. Тесты - управление тестированием", None),
            ("", None),
            ("РАЗДЕЛ III: ГРАФИКИ И ВИЗУАЛИЗАЦИЯ", "Heading 2"),
            ("    9. Графики производительности менеджеров", None),
            ("    10. Динамика оценок качества", None),
            ("    11. Статус прохождения тестов", None),
            ("    12. Динамика тестирования", None),
            ("", None),
            ("РАЗДЕЛ IV: НАСТРОЙКА АНАЛИТИЧЕСКИХ ТАБЛИЦ", "Heading 2"),
            ("    13. Создание кастомных таблиц", None),
            ("    14. Работа с ИИ-промтами", None),
            ("    15. Готовые примеры промтов", None),
        ]
        
        for item_text, style in toc_content:
            if not item_text:
                doc.add_paragraph()
                continue
            
            if style:
                p = doc.add_paragraph(item_text, style=style)
            else:
                p = doc.add_paragraph(item_text)
                p.paragraph_format.left_indent = Inches(0.25)
        
        doc.add_page_break()
    
    def find_screenshot(self, filename):
        """Поиск скриншота используя JSON маппинг и fallback к папкам"""
        filename = filename.strip()
        
        # Убираем все звездочки из имени файла
        filename = filename.replace('*', '').strip()
        
        # Если нет расширения, добавляем .png
        if not filename.endswith('.png'):
            filename += '.png'
        
        # Сначала проверяем в JSON маппинге
        if filename in self.screenshot_mapping:
            screenshot_path = self.base_path / self.screenshot_mapping[filename]
            if screenshot_path.exists():
                return screenshot_path
        
        # Fallback: поиск по частичному совпадению в маппинге
        filename_base = filename.lower().replace('.png', '').replace('.jpg', '')
        for mapped_name, mapped_path in self.screenshot_mapping.items():
            mapped_base = mapped_name.lower().replace('.png', '').replace('.jpg', '')
            if filename_base in mapped_base or mapped_base in filename_base:
                screenshot_path = self.base_path / mapped_path
                if screenshot_path.exists():
                    return screenshot_path
        
        # Fallback: поиск во всех подпапках (старый метод)
        for root, dirs, files in os.walk(self.screenshots_path):
            for file in files:
                # Точное совпадение
                if file.lower() == filename.lower():
                    return Path(root) / file
                
                # Частичное совпадение (убираем расширения для сравнения)
                file_base = file.lower().replace('.png', '').replace('.jpg', '')
                filename_base = filename.lower().replace('.png', '').replace('.jpg', '')
                
                if filename_base in file_base or file_base in filename_base:
                    return Path(root) / file
        return None
    
    def process_text_formatting(self, text):
        """Очищает markdown символы кроме важных для форматирования"""
        # Проверяем есть ли метки перед обработкой
        if '{INTERFACE}' in text or '{TECHNICAL}' in text:
            print(f"DEBUG: Found markup in process_text_formatting: {text}")
        
        # Убираем только лишние markdown символы, оставляем ** для жирного текста
        text = re.sub(r'[#`]', '', text).strip()
        return text
    
    def add_formatted_paragraph(self, doc, text, style=None):
        """Добавляет параграф с правильным форматированием жирного текста и гиперссылок"""
        if style:
            p = doc.add_paragraph(style=style)
        else:
            p = doc.add_paragraph()
        
        # Проверяем, есть ли гиперссылки
        if '[' in text and '](#' in text:
            # Обрабатываем текст с гиперссылками и жирным текстом
            self._add_complex_formatting(p, text)
        else:
            # Обрабатываем только жирный текст **text**
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Жирный текст
                    bold_text = part[2:-2]  # Убираем **
                    run = p.add_run(bold_text)
                    run.font.bold = True
                elif part:
                    # Обычный текст
                    clean_part = re.sub(r'[*#`]', '', part)
                    if clean_part.strip():
                        run = p.add_run(clean_part)
        
        return p
    
    def _add_complex_formatting(self, paragraph, text):
        """Добавляет текст с комбинированным форматированием (жирный + гиперссылки)"""
        import re
        
        # Сначала обрабатываем гиперссылки, затем жирный текст в каждой части
        hyperlink_pattern = r'\[([^\]]+)\]\(#([^)]+)\)'
        parts = re.split(hyperlink_pattern, text)
        
        for i, part in enumerate(parts):
            if i % 3 == 0:  # Обычный текст (может содержать жирный)
                if part:
                    self._add_bold_formatting(paragraph, part)
            elif i % 3 == 1:  # Текст ссылки
                link_text = part
            elif i % 3 == 2:  # URL ссылки
                anchor = "#" + part
                target_heading = self.hyperlink_mapping.get(anchor, None)
                
                if target_heading:
                    self._create_internal_hyperlink(paragraph, link_text, target_heading)
                else:
                    hyperlink = paragraph.add_run(link_text)
                    hyperlink.font.color.rgb = RGBColor(255, 0, 0)
    
    def _add_bold_formatting(self, paragraph, text):
        """Добавляет текст с поддержкой жирного форматирования"""
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                bold_text = part[2:-2]
                run = paragraph.add_run(bold_text)
                run.font.bold = True
            elif part:
                clean_part = re.sub(r'[*#`]', '', part)
                if clean_part.strip():
                    run = paragraph.add_run(clean_part)
    
    def add_interface_block(self, doc, text):
        """Добавляет блок описания интерфейса с полной левой границей"""
        print(f"DEBUG: Creating interface block with full border: {text[:50]}...")
        
        # Создаем таблицу 1x1 для полной границы
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        
        # Настраиваем границы таблицы
        from docx.oxml import parse_xml
        from docx.oxml.shared import qn
        
        # Только левая граница синего цвета
        borders_xml = '''
        <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:left w:val="single" w:sz="24" w:color="3498DB"/>
            <w:top w:val="nil"/>
            <w:right w:val="nil"/>  
            <w:bottom w:val="nil"/>
        </w:tcBorders>'''
        
        # Светло-голубой фон
        shading_xml = '''
        <w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
               w:val="clear" w:color="auto" w:fill="EBF3FD"/>'''
        
        try:
            # Применяем границы и фон
            borders_element = parse_xml(borders_xml)
            shading_element = parse_xml(shading_xml)
            
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_pr.append(borders_element)
            tc_pr.append(shading_element)
        except Exception as e:
            print(f"DEBUG: Border/shading failed: {e}")
        
        # Добавляем текст в ячейку
        cell_para = cell.paragraphs[0]
        self._add_formatted_text_to_paragraph(cell_para, text)
        
        # Форматирование текста в ячейке
        cell_para.paragraph_format.space_before = Pt(6)
        cell_para.paragraph_format.space_after = Pt(6)
        
        for run in cell_para.runs:
            run.font.color.rgb = RGBColor(40, 40, 40)
            run.font.italic = True
        
        return table
    
    def _add_formatted_text_to_paragraph(self, paragraph, text):
        """Добавляет форматированный текст в параграф с поддержкой гиперссылок и жирного текста"""
        # Проверяем есть ли гиперссылки
        if '[' in text and '](#' in text:
            # Используем сложное форматирование с гиперссылками
            self._add_complex_formatting(paragraph, text)
        else:
            # Обрабатываем только жирный текст **text**
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Жирный текст
                    bold_text = part[2:-2]  # Убираем **
                    run = paragraph.add_run(bold_text)
                    run.font.bold = True
                elif part:
                    # Обычный текст
                    clean_part = re.sub(r'[*#`]', '', part)
                    if clean_part.strip():
                        run = paragraph.add_run(clean_part)
    
    def add_technical_block(self, doc, text):
        """Добавляет технический блок с полной оранжевой границей"""
        print(f"DEBUG: Creating technical block with full border: {text[:50]}...")
        
        # Создаем таблицу 1x1 для полной границы
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        
        # Настраиваем границы таблицы
        from docx.oxml import parse_xml
        
        # Только левая граница оранжевого цвета
        borders_xml = '''
        <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:left w:val="single" w:sz="24" w:color="E67E22"/>
            <w:top w:val="nil"/>
            <w:right w:val="nil"/>  
            <w:bottom w:val="nil"/>
        </w:tcBorders>'''
        
        # Светло-оранжевый фон
        shading_xml = '''
        <w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
               w:val="clear" w:color="auto" w:fill="FDF2E9"/>'''
        
        try:
            # Применяем границы и фон
            borders_element = parse_xml(borders_xml)
            shading_element = parse_xml(shading_xml)
            
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_pr.append(borders_element)
            tc_pr.append(shading_element)
        except Exception as e:
            print(f"DEBUG: Technical border/shading failed: {e}")
        
        # Добавляем текст в ячейку
        cell_para = cell.paragraphs[0]
        self._add_formatted_text_to_paragraph(cell_para, text)
        
        # Форматирование технического текста
        cell_para.paragraph_format.space_before = Pt(6)
        cell_para.paragraph_format.space_after = Pt(6)
        
        for run in cell_para.runs:
            run.font.color.rgb = RGBColor(80, 80, 80)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        
        return table
    
    def add_code_block(self, doc, text, title=""):
        """Добавляет красивый блок кода с фоном"""
        # Создаем таблицу для имитации блока кода
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        
        cell = table.cell(0, 0)
        cell.width = Inches(6)
        
        # Фон ячейки (имитация цветного фона)
        from docx.oxml.shared import qn
        from docx.oxml import parse_xml
        
        shading_elm = parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Добавляем заголовок если есть
        if title:
            title_p = cell.paragraphs[0]
            title_run = title_p.add_run(title)
            title_run.font.bold = True
            title_run.font.size = Pt(10)
            title_run.font.color.rgb = RGBColor(237, 28, 36)  # Setl Group красный
            cell.add_paragraph()
        
        # Добавляем код
        code_p = cell.add_paragraph()
        code_run = code_p.add_run(text)
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(9)
        code_run.font.color.rgb = RGBColor(40, 40, 40)
        
        return table
    
    def add_hyperlink(self, paragraph, text):
        """Добавляет текст с гиперссылками в параграф"""
        import re
        from docx.oxml.shared import qn
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        # Ищем паттерн [текст](#ссылка)
        hyperlink_pattern = r'\[([^\]]+)\]\(#([^)]+)\)'
        
        # Разбиваем текст на части
        parts = re.split(hyperlink_pattern, text)
        
        for i, part in enumerate(parts):
            if i % 3 == 0:  # Обычный текст
                if part:
                    self._add_bold_formatting(paragraph, part)
            elif i % 3 == 1:  # Текст ссылки
                link_text = part
            elif i % 3 == 2:  # URL ссылки
                anchor = "#" + part
                target_heading = self.hyperlink_mapping.get(anchor, None)
                
                if target_heading:
                    # Создаем реальную внутреннюю гиперссылку в Word
                    self._create_internal_hyperlink(paragraph, link_text, target_heading)
                    print(f"DEBUG: Создана гиперссылка '{link_text}' -> '{target_heading}'")
                else:
                    # Если маппинг не найден, показываем красным
                    hyperlink = paragraph.add_run(link_text)
                    hyperlink.font.color.rgb = RGBColor(255, 0, 0)
                    print(f"WARNING: Не найден маппинг для гиперссылки: {anchor}")
    
    def _create_internal_hyperlink(self, paragraph, link_text, target_heading):
        """Создает внутреннюю гиперссылку Word"""
        from docx.oxml.shared import qn, OxmlElement
        
        # Создаем ID для закладки на основе заголовка
        bookmark_id = target_heading.replace(" ", "_").replace(".", "_").replace("-", "_").replace("(", "").replace(")", "")
        
        # Создаем элемент гиперссылки
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_id)
        
        # Создаем run внутри гиперссылки
        new_run = OxmlElement('w:r')
        
        # Создаем свойства run с форматированием ссылки
        rPr = OxmlElement('w:rPr')
        
        # Добавляем синий цвет
        c = OxmlElement('w:color')
        c.set(qn('w:val'), '2A6099')
        rPr.append(c)
        
        # Добавляем подчеркивание
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        
        # Добавляем текст ссылки
        t = OxmlElement('w:t')
        t.text = link_text
        new_run.append(t)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        
        print(f"DEBUG: Создана настоящая гиперссылка '{link_text}' -> закладка '{bookmark_id}'")
    
    def _add_bookmark(self, paragraph, bookmark_name):
        """Добавляет закладку к параграфу"""
        from docx.oxml.shared import qn, OxmlElement
        
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        r = run._r
        
        # Создаем начало закладки
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), '0')
        start.set(qn('w:name'), bookmark_name)
        r.insert(0, start)
        
        # Создаем конец закладки
        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), '0')
        end.set(qn('w:name'), bookmark_name)
        r.append(end)
        
        print(f"DEBUG: Создана закладка '{bookmark_name}'")
    
    def process_text_with_hyperlinks(self, doc, text):
        """Обрабатывает текст с гиперссылками и добавляет в документ"""
        import re
        
        # Проверяем, есть ли в тексте гиперссылки
        if '[' in text and '](#' in text:
            p = doc.add_paragraph()
            self.add_hyperlink(p, text)
        else:
            p = doc.add_paragraph(text)
        
        return p
    
    def read_part_content(self, part_name):
        """Чтение содержимого части инструкции"""
        # Сначала проверяем версию _NEW если существует
        new_file = self.parts_path / f"{part_name}_NEW.md"
        if new_file.exists():
            with open(new_file, 'r', encoding='utf-8') as f:
                return f.read()
        
        # Иначе используем обычный файл
        part_file = self.parts_path / f"{part_name}.md"
        if part_file.exists():
            with open(part_file, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            return f"Данный раздел находится в разработке и будет добавлен в следующих версиях системы."
    
    def add_section_to_doc(self, doc, section_key, force_regenerate=False):
        """Добавление раздела в документ"""
        section_info = self.sections[section_key]
        
        # Проверяем, нужно ли генерировать раздел
        if not force_regenerate and section_key in self.progress['completed_sections']:
            print(f"Раздел '{section_info['title']}' уже создан, пропускаем...")
            return
        
        print(f"Создаём раздел: {section_info['title']}")
        
        # Заголовок раздела
        section_heading = doc.add_paragraph(section_info['title'], style='Heading 1')
        
        # Добавляем закладку для заголовка раздела
        section_bookmark_id = section_info['title'].replace(" ", "_").replace(".", "_").replace("-", "_").replace("(", "").replace(")", "")
        self._add_bookmark(section_heading, section_bookmark_id)
        
        # Части раздела
        for part_name in section_info['parts']:
            if not force_regenerate and part_name in self.progress['completed_parts']:
                print(f"  Часть '{part_name}' уже создана, пропускаем...")
                continue
                
            print(f"  Добавляем часть: {part_name}")
            
            # Читаем содержимое части
            content = self.read_part_content(part_name)
            print(f"DEBUG: Content preview for {part_name}: {content[:100]}...")
            print(f"DEBUG: File used for {part_name}: {self.parts_path / f'{part_name}.md'}")
            print(f"DEBUG: Has INTERFACE tags: {'{INTERFACE}' in content}")
            
            # Профессиональная обработка markdown -> docx
            lines = content.split('\n')
            current_list_level = 0
            
            in_code_block = False
            code_block_content = []
            
            for line_num, line in enumerate(lines, 1):
                line = line.rstrip()
                
                # Отладка для поиска меток
                if '{INTERFACE}' in line or '{TECHNICAL}' in line:
                    print(f"DEBUG: Line {line_num}: {line}")
                
                # СНАЧАЛА проверяем специальные блоки (высший приоритет)
                if line.startswith('{INTERFACE}'):
                    text = line[11:].strip()  # Убираем метку {INTERFACE}
                    print(f"DEBUG: Processing INTERFACE block: {text}")
                    self.add_interface_block(doc, text)
                    current_list_level = 0
                    continue
                elif line.startswith('{TECHNICAL}'):
                    text = line[11:].strip()  # Убираем метку {TECHNICAL}
                    print(f"DEBUG: Processing TECHNICAL block: {text}")
                    self.add_technical_block(doc, text)
                    current_list_level = 0
                    continue
                
                # Обработка блоков кода ```
                if line.strip() == '```':
                    if in_code_block:
                        # Конец блока кода
                        if code_block_content:
                            self.add_code_block(doc, '\\n'.join(code_block_content), "API Endpoint:")
                        code_block_content = []
                        in_code_block = False
                    else:
                        # Начало блока кода
                        in_code_block = True
                    continue
                
                if in_code_block:
                    code_block_content.append(line)
                    continue
                
                if not line:
                    # Пустая строка
                    doc.add_paragraph()
                    current_list_level = 0
                    continue
                
                if line.startswith('# '):
                    # Заголовок части (Heading 2) - убираем решетку
                    clean_title = line[2:].strip()
                    heading_para = doc.add_paragraph(clean_title, style='Heading 2')
                    
                    # Добавляем закладку для заголовка
                    bookmark_id = clean_title.replace(" ", "_").replace(".", "_").replace("-", "_").replace("(", "").replace(")", "")
                    self._add_bookmark(heading_para, bookmark_id)
                    current_list_level = 0
                    
                elif line.startswith('## '):
                    # Подзаголовок (Heading 3) - убираем решетки
                    clean_title = line[3:].strip()
                    doc.add_paragraph(clean_title, style='Heading 3')
                    current_list_level = 0
                    
                elif line.startswith('### '):
                    # Подподзаголовок - убираем решетки
                    clean_title = line[4:].strip()
                    p = doc.add_paragraph(clean_title)
                    p.runs[0].font.bold = True
                    p.runs[0].font.size = Pt(12)
                    p.runs[0].font.color.rgb = RGBColor(122, 122, 122)  # Setl Group серый
                    current_list_level = 0
                    
                elif line.startswith('|') and '|' in line.strip():
                    # Обработка таблиц Markdown
                    table_lines = [line]
                    # Собираем ВСЕ строки таблицы до первой НЕ-табличной строки
                    for next_line_idx in range(line_num + 1, len(lines)):
                        next_line = lines[next_line_idx].strip()
                        
                        # Останавливаемся на заголовке или тексте (НЕ на пустой строке!)
                        if next_line.startswith('#') or (next_line and not next_line.startswith('|')):
                            break
                            
                        # Добавляем только строки с таблицей (игнорируем пустые)
                        if next_line.startswith('|') and '|' in next_line:
                            table_lines.append(next_line)
                    
                    # Создаем таблицу в Word
                    if len(table_lines) >= 2:  # Минимум заголовок и одна строка
                        self.add_markdown_table(doc, table_lines)
                        # Пропускаем обработанные строки таблицы
                        continue
                    current_list_level = 0
                    
                elif '.png' in line or 'скриншот' in line.lower() or '[СКРИНШОТ:' in line:
                    # Обработка разных форматов скриншотов
                    if '[СКРИНШОТ:' in line:
                        # Новый формат [СКРИНШОТ: имя.png]
                        match = re.search(r'\[СКРИНШОТ:\s*([^\]]+)\]', line)
                        if match:
                            image_filename = match.group(1).strip()
                            screenshot_text = image_filename
                        else:
                            continue
                    else:
                        # Старый формат
                        clean_line = line.replace('*', '').strip()
                        
                        # Извлекаем имя файла и описание
                        if ' - ' in clean_line:
                            parts = clean_line.split(' - ', 1)
                            image_filename = parts[0].strip()
                            description = parts[1].strip()
                            screenshot_text = f"{image_filename} - {description}"
                        else:
                            image_filename = clean_line.strip()
                            screenshot_text = image_filename
                    
                    # Ищем изображение
                    image_path = self.find_screenshot(image_filename)
                    if image_path and image_path.exists():
                        # Вставляем изображение
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(str(image_path), width=Inches(6))
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Добавляем подпись под изображением
                        caption = doc.add_paragraph(screenshot_text, style='Screenshot Reference')
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        # Если изображение не найдено, просто пропускаем без сообщения
                        pass
                    current_list_level = 0
                    
                elif line.startswith('URL') and ':' in line:
                    # URL без звездочек
                    text = line.replace('URL для входа', 'URL для входа')
                    if any(keyword in text.lower() for keyword in ['важно', 'внимание', 'примечание', 'url']):
                        doc.add_paragraph(text, style='Important Note')
                    else:
                        p = doc.add_paragraph(text)
                        p.runs[0].font.bold = True
                    current_list_level = 0
                    
                elif line.startswith('- ') or line.startswith('* '):
                    # Обычный список с форматированием
                    list_text = line[2:].strip()
                    if list_text:
                        p = self.add_formatted_paragraph(doc, list_text, 'List Bullet')
                    current_list_level = 1
                    
                elif re.match(r'^\d+\.\s', line):
                    # Нумерованный список - убираем markdown
                    list_text = re.sub(r'^\d+\.\s', '', line)
                    list_text = re.sub(r'[*#`]', '', list_text).strip()
                    if list_text:
                        p = doc.add_paragraph(list_text, style='List Number')
                    current_list_level = 1
                    
                elif line.startswith('  - ') or line.startswith('    - '):
                    # Вложенный список
                    indent_level = (len(line) - len(line.lstrip())) // 2
                    list_text = line.strip()[2:]
                    p = doc.add_paragraph(list_text, style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.25 * (indent_level // 2 + 1))
                    current_list_level = 2
                    
                elif line.startswith('http://') or line.startswith('https://'):
                    # URL в красивом блоке
                    self.add_code_block(doc, line, "URL:")
                    current_list_level = 0
                    
                elif '`' in line:
                    # Строка с кодом - убираем все звездочки
                    clean_line = line.replace('*', '')
                    p = doc.add_paragraph()
                    parts = clean_line.split('`')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            # Обычный текст
                            if part:
                                p.add_run(part)
                        else:
                            # Код
                            run = p.add_run(part)
                            run.style = 'Code Text'
                    current_list_level = 0
                    
                elif line.strip():
                    # Обрабатываем специальные блоки и обычный текст
                    if line.startswith('{INTERFACE}'):
                        # Блок описания интерфейса с левой границей
                        print(f"DEBUG: Found INTERFACE block: {line}")
                        text = line[11:].strip()  # Убираем метку {INTERFACE}
                        print(f"DEBUG: Interface text: {text}")
                        self.add_interface_block(doc, text)
                    elif line.startswith('{TECHNICAL}'):
                        # Технический блок с синей границей
                        print(f"DEBUG: Found TECHNICAL block: {line}")
                        text = line[11:].strip()  # Убираем метку {TECHNICAL}
                        self.add_technical_block(doc, text)
                    else:
                        # Обычный текст - обрабатываем жирный текст и убираем остальное markdown
                        clean_line = self.process_text_formatting(line.strip())
                        if clean_line:
                            # Проверяем на важные примечания
                            if any(keyword in clean_line.lower() for keyword in ['важно', 'внимание', 'примечание', 'note']):
                                p = self.add_formatted_paragraph(doc, line, 'Important Note')
                            else:
                                p = self.add_formatted_paragraph(doc, line)
                                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    current_list_level = 0
            
            # Отмечаем часть как завершённую
            if part_name not in self.progress['completed_parts']:
                self.progress['completed_parts'].append(part_name)
        
        # Отмечаем раздел как завершённый
        if section_key not in self.progress['completed_sections']:
            self.progress['completed_sections'].append(section_key)
        
        # Разрыв страницы после раздела
        doc.add_page_break()
    
    def generate(self, sections_to_generate=None, force_regenerate=False):
        """Основная функция генерации"""
        print("Запуск генерации инструкции...")
        
        # Создаём или открываем документ
        if self.output_path.exists() and not force_regenerate:
            print("Открываем существующий документ...")
            doc = Document(str(self.output_path))
        else:
            print("Создаём новый документ...")
            doc = self.create_document()
            self.add_table_of_contents(doc)
            # Очищаем прогресс при создании нового документа
            if force_regenerate:
                self.progress = {"completed_sections": [], "completed_parts": []}
        
        # Определяем какие разделы генерировать
        if sections_to_generate:
            target_sections = sections_to_generate
        else:
            target_sections = list(self.sections.keys())
        
        # Генерируем разделы
        for section_key in target_sections:
            if section_key in self.sections:
                self.add_section_to_doc(doc, section_key, force_regenerate)
            else:
                print(f"ОШИБКА: Неизвестный раздел: {section_key}")
        
        # Сохраняем документ и прогресс
        print("Сохраняем документ...")
        doc.save(str(self.output_path))
        self.save_progress()
        
        print(f"ГОТОВО: Инструкция сохранена: {self.output_path}")
        print(f"Прогресс: {len(self.progress['completed_parts'])} частей создано")
    
    def add_markdown_table(self, doc, table_lines):
        """Добавляет таблицу Markdown в документ Word"""
        from docx.shared import RGBColor, Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        # Фильтруем пустые строки и строки-разделители
        valid_lines = []
        for line in table_lines:
            line = line.strip()
            if line and line.startswith('|') and not line.replace('|', '').replace('-', '').replace(' ', '').replace(':', '') == '':
                valid_lines.append(line)
        
        if len(valid_lines) < 2:  # Нужен минимум заголовок + строка данных
            return
        
        # Парсим заголовок
        header_line = valid_lines[0]
        headers = [cell.strip() for cell in header_line.split('|')[1:-1]]  # Убираем пустые первый и последний элементы
        
        # Парсим строки данных (пропускаем разделитель если есть)
        data_lines = []
        for line in valid_lines[1:]:
            if not line.replace('|', '').replace('-', '').replace(' ', '').replace(':', '') == '':
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if len(cells) == len(headers):  # Проверяем соответствие количества колонок
                    data_lines.append(cells)
        
        if not data_lines:
            return
        
        # Создаем таблицу в Word
        table = doc.add_table(rows=1 + len(data_lines), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = 'Table Grid'
        
        # Заполняем заголовок
        header_row = table.rows[0]
        for i, header_text in enumerate(headers):
            cell = header_row.cells[i]
            paragraph = cell.paragraphs[0]
            paragraph.text = header_text
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Форматирование заголовка
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)  # Белый текст
            
            # Фон заголовка
            from docx.oxml import parse_xml
            shading_elm = parse_xml(r'<w:shd {} w:fill="2C3E50"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Заполняем данные
        for row_idx, row_data in enumerate(data_lines, 1):
            row = table.rows[row_idx]
            for col_idx, cell_text in enumerate(row_data):
                cell = row.cells[col_idx]
                paragraph = cell.paragraphs[0]
                paragraph.text = cell_text
                
                # Форматирование данных
                run = paragraph.runs[0]
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(60, 60, 60)
        
        print(f"DEBUG: Создана таблица {len(headers)}x{len(data_lines)+1}")

def main():
    parser = argparse.ArgumentParser(description='Генератор инструкции Цифровой РОП')
    parser.add_argument('--sections', nargs='+', help='Конкретные разделы для генерации')
    parser.add_argument('--force', action='store_true', help='Принудительная перезапись всех разделов')
    parser.add_argument('--reset', action='store_true', help='Сброс прогресса и создание нового документа')
    parser.add_argument('--list', action='store_true', help='Показать доступные разделы')
    
    args = parser.parse_args()
    
    # Определяем базовый путь
    base_path = Path(__file__).parent.parent
    generator = InstructionGenerator(base_path)
    
    if args.list:
        print("Доступные разделы:")
        for key, info in generator.sections.items():
            print(f"  {key}: {info['title']}")
        return
    
    if args.reset:
        print("Сброс прогресса...")
        if generator.progress_file.exists():
            generator.progress_file.unlink()
        if generator.output_path.exists():
            generator.output_path.unlink()
        generator.load_progress()
    
    # Запуск генерации
    generator.generate(
        sections_to_generate=args.sections,
        force_regenerate=args.force or args.reset
    )

if __name__ == "__main__":
    main()